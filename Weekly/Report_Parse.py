# -*- coding: utf-8 -*-
"""
Created on Thu Jan 23 21:58:26 2020

@author: LiJinling
"""

import docx
import openpyxl
import re
import os
import collections

def word_parse(file_name,wb):
    
    file=docx.Document(file_name)
    title_dic = collections.OrderedDict()   
    content_dic={}
    para_list=[]
    title_list=[]
    para_no=0
    #获取标题段落位置
    for para in file.paragraphs:
        if len(para.text) < 1:
            file.paragraphs[para_no].clear()
            para_no-=1
        else:
            para_list.append(para.text)
        if '签发：' in para.text:
            start_time=re.split('－| ',para.text)[0]
            end_time=re.split('－| ',para.text)[1]
        elif '期）' in para.text and '第'in para.text and'年'in para.text:
            week_no=re.split('（|）',para.text)[1]
        elif '【部门总体情况】'in para.text:
            title_dic['【部门总体情况】']=para_no
            title_list.append(['【部门总体情况】',para_no])
        elif '【部门重要工作】'in para.text:
            title_dic['【部门重要工作】']=para_no
            title_list.append(['【部门重要工作】',para_no])
        elif '【项目进展】'in para.text:
            title_dic['【项目进展】']=para_no
            title_list.append(['【项目进展】',para_no])
        elif '【重点需求'in para.text:
            title_dic['【重点需求】']=para_no
            title_list.append(['【重点需求】',para_no])
        elif '【部门管理】'in para.text:
            title_dic['【部门管理】']=para_no
            title_list.append(['【部门管理】',para_no])
        elif '【系统运行情况】'in para.text:
            title_dic['【系统运行情况】']=para_no
            title_list.append(['【系统运行情况】',para_no])
        elif '【党建工作】'in para.text:
            title_dic['【党建工作】']=para_no
            title_list.append(['【党建工作】',para_no])
        elif '【存在问题及建议】'in para.text:
            title_dic['【存在问题及建议】']=para_no
            title_list.append(['【存在问题及建议】',para_no])
        para_no+=1
    
    for i in range(len(title_list)):
        if(i<len(title_list)-1):
            content_dic[title_list[i][0]]=para_list[title_list[i][1]+1:title_list[i+1][1]]
        else:
            content_dic[title_list[i][0]]=para_list[title_list[i][1]+1:len(para_list)-2]
        
    for key in content_dic:
        for k in content_dic[key]:
            wb.get_sheet_by_name(key).append([week_no,start_time,end_time,k])

cwd=os.getcwd()
wb = openpyxl.Workbook()
for files in os.walk(cwd):
    file_list=files[2]  


wb.create_sheet('【部门总体情况】').append(['期数','开始日期','结束日期','内容'])
wb.create_sheet('【部门重要工作】').append(['期数','开始日期','结束日期','内容'])
wb.create_sheet('【项目进展】').append(['期数','开始日期','结束日期','内容'])
wb.create_sheet('【重点需求】').append(['期数','开始日期','结束日期','内容'])
wb.create_sheet('【部门管理】').append(['期数','开始日期','结束日期','内容'])
wb.create_sheet('【系统运行情况】').append(['期数','开始日期','结束日期','内容'])
wb.create_sheet('【党建工作】').append(['期数','开始日期','结束日期','内容'])
wb.create_sheet('【存在问题及建议】').append(['期数','开始日期','结束日期','内容'])
wb.remove_sheet(wb["Sheet"])
   
for file_name in file_list:
    if '信息科技部周报' in file_name and 'doc'in file_name:
        print(file_name)
        word_parse(file_name,wb)    
wb.save('周报.xlsx')
        

    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    