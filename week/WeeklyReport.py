# -*- coding: utf-8 -*-
"""
Created on Thu Jan 23 14:55:07 2020

@author: LiJinling
"""
import docx
import openpyxl
import re

#获取文档对象

file=docx.Document("信息科技部周报2020年第3期（2020年1月13日-2019年1月19日）.docx")
#创建一个工作薄
wb = openpyxl.Workbook()
#保存文件
wb.save('周报.xlsx')
#新增一个sheet表单
wb.create_sheet('部门总体情况')

para_no = 0
para_num=len(file.paragraphs)
#for para in file.paragraphs:
while para_no<para_num:
#    print('第',parag_num,'段')
    para_text=file.paragraphs[para_no].text
    #print(para_text)
    #获取周报日期
    if '签发：' in para_text:
        start_time=re.split('－| ',para_text)[0]
        end_time=re.split('－| ',para_text)[1]
        print('start_time=',start_time)
        print('end_time=',end_time)
        para_no +=1
        continue
    #获取部门总体情况
    if '【部门总体情况】' in para_text:
        print( '【部门总体情况】')
        all_content=file.paragraphs[para_no+1].text
        print(all_content)
        para_no +=1
        continue
    #获取部门重点工作
    if '【部门重要工作】' in para_text:
        print( '【部门重要工作】')
        para_no +=1
        while True:       
            content=file.paragraphs[para_no].text
            print(content)
            para_no +=1           
            if('【' in file.paragraphs[para_no].text) or para_no>=para_num:break
    #获取【项目进展】
    if '【项目进展】' in para_text:
        print('【项目进展】')
        para_no +=1
        while True:       
            content=file.paragraphs[para_no].text
            print(content)
            para_no +=1            
            if('【' in file.paragraphs[para_no].text) or para_no>=para_num:break
            #continue
    if '【重点需求】' in para_text:
        print( '【重点需求】')
        para_no +=1
        while True:       
            content=file.paragraphs[para_no].text
            print(content)
            para_no +=1            
            if('【' in file.paragraphs[para_no].text) or para_no>=para_num:break
            
            #continue
    #获取【部门管理】
    if '【部门管理】' in para_text:
        print( '【部门管理】')
        para_no +=1
        while True:       
            content=file.paragraphs[para_no].text
            print(content)
            para_no +=1            
            if('【' in file.paragraphs[para_no].text) or para_no>=para_num:break
            #continue
        #【系统运行情况】
    if '【系统运行情况】' in para_text:
        print( '【系统运行情况】')
        para_no +=1
        while True:       
            content=file.paragraphs[para_no].text
            print(content)
            para_no +=1            
            if('【' in file.paragraphs[para_no].text) or para_no>=para_num:break
                #continue
    #【党建工作】
    if '【党建工作】' in para_text:
        print( '【党建工作】')
        para_no +=1
        while True:       
            content=file.paragraphs[para_no].text
            print(content)
            para_no +=1            
            if('【' in file.paragraphs[para_no].text) or para_no>=para_num:break
            #continue    
#    if '【存在问题及建议】' in para_text:
#        print( '【存在问题及建议】')
#        para_no +=1
#        while True:       
#            content=file.paragraphs[para_no].text
#            print(content)
#            para_no +=1            
#            if('【' in file.paragraphs[para_no].text) or para_no>=para_num:break
            #continue   

    para_no +=1
    #print('para_no=',para_no)
#    parag_num += 1  
#print ('This document has ', parag_num, ' paragraphs')